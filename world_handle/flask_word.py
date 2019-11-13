# 需求，根据项目页面展示数据生成对应的文档
# 然后flask 接口下载生成的word文档
# -*- coding: utf-8 -*-
import re
import io
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from flask import send_file

from tests.get_docx import get_person_info, get_org_info, get_event_info, get_country_info
from urllib.parse import quote


# 基本信息
def judge_text_or_img(document, tittle, info):
    p = document.add_paragraph()
    run = p.add_run(tittle)
    run.font.name = u'仿宋'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    run.font.size = Pt(16)

    get_p = re.compile("<p>.*</p>")
    results_p = get_p.findall(info)
    # print(results_p)
    for result_p in results_p:
        if result_p.find("<img alt") == -1:
            # 文字 做文字处理
            # print(result_p)
            end_p = result_p.find("</p>")
            text_p = result_p[3:end_p]
            p = document.add_paragraph()
            run = p.add_run("{}".format(text_p.replace("&nbsp;", " ")))
            run.font.name = u'仿宋'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            run.font.size = Pt(16)
            # print(text_p)
        else:
            # 图片做图片处理
            start_img = result_p.find("src=") + 5
            end_img = result_p.find("style=") - 2
            img_url = result_p[start_img:end_img]

            img_url_new = '/usr/share/nginx/html/sea_survival/storage/app/public/' + img_url[35:]

            # print(result_p)
            print("---------------------------------------------")
            print(img_url)
            print(img_url_new)
            print("---------------------------------------------")
            document.add_picture('{}'.format(img_url_new), width=Inches(1.25))

            # print(result_p.find("src="))
            # print(result_p.find("style="))


# 大类下边直接内容
def get_big_type(document, info):
    p = document.add_paragraph()
    get_p = re.compile("<p>.*</p>")
    results_p = get_p.findall(info)
    # print(results_p)
    for result_p in results_p:
        if result_p.find("<img alt") == -1:
            # 文字 做文字处理
            # print(result_p)
            end_p = result_p.find("</p>")
            text_p = result_p[3:end_p]
            p = document.add_paragraph()
            run = p.add_run("{}".format(text_p.replace("&nbsp;", " ")))
            run.font.name = u'仿宋'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
            run.font.size = Pt(16)
            # print(text_p)
        else:
            # 图片做图片处理
            start_img = result_p.find("src=") + 5
            end_img = result_p.find("style=") - 2
            img_url = result_p[start_img:end_img]

            # img_url_new = 'D:\\xyh_project\\storage\\' + img_url[35:]
            img_url_new = '/usr/share/nginx/html/sea_survival/storage/app/public/' + img_url[35:]
            print(img_url)
            print(img_url_new)
            document.add_picture('{}'.format(img_url_new), width=Inches(2.25))


# 小标题 比如简介
def get_little_tittle(document, title):
    p = document.add_paragraph()
    run = p.add_run("{}".format(title))
    run.font.name = u'宋体（正文）'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体（正文）')
    run.font.size = Pt(22)
    run.bold = True


# 简单字段
def get_same_format(document, tittle, info):
    p = document.add_paragraph()
    run = p.add_run("{}{}".format(tittle, info))
    run.font.name = u'仿宋'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    run.font.size = Pt(16)


def get_person_docx(id):
    results_dict = get_person_info(id)

    person_name = results_dict['person_name']
    other_name = results_dict['other_name']
    header_img = results_dict['header_img']
    sex = results_dict['sex']
    birthday = results_dict['birthday']
    birth_address = results_dict['birth_address']
    country = results_dict['country']
    motto = results_dict['motto']
    faith = results_dict['faith']
    partisan = results_dict['partisan']
    career = results_dict['career']
    language = results_dict['language']
    id_number = results_dict['id_number']
    ed_experience = results_dict['ed_experience']
    work_experience = results_dict['work_experience']
    unit = results_dict['unit']
    characterisics = results_dict['characterisics']
    economic_status = results_dict['economic_status']
    polotical_views = results_dict['polotical_views']
    criminal_record = results_dict['criminal_record']
    controversial = results_dict['controversial']
    import_event = results_dict['import_event']
    main_achieve = results_dict['main_achieve']
    family_member = results_dict['family_member']
    social_relationship = results_dict['social_relationship']
    contact_information = results_dict['contact_information']
    video_material = results_dict['video_material']
    others = results_dict['others']

    # 主文档
    document = Document()

    # 标题
    p = document.add_paragraph()
    run = p.add_run(person_name + "调研")
    run.font.name = u'宋体（正文）'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体（正文）')
    run.font.size = Pt(22)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中
    run.bold = True  # 加粗

    # 简介
    get_little_tittle(document, "一、简介")

    # 姓名
    tittle = "姓名："
    info = person_name
    get_same_format(document, tittle, info)

    # 别名
    tittle = "别名："
    info = other_name
    get_same_format(document, tittle, info)

    # 头像
    tittle = "头像："
    info = header_img
    judge_text_or_img(document, tittle, header_img)

    # 性别
    tittle = "性别："
    info = sex
    get_same_format(document, tittle, info)

    # 生日
    tittle = "生日："
    info = birthday
    get_same_format(document, tittle, info)

    # 出生地
    tittle = "出生地："
    info = birth_address
    get_same_format(document, tittle, info)

    # 国籍
    tittle = "国籍："
    info = country
    get_same_format(document, tittle, info)

    # 座右铭
    tittle = "座右铭："
    judge_text_or_img(document, tittle, motto)
    # print(person_name)
    # print(motto)
    # print(polotical_views)

    # 宗教信仰
    tittle = "宗教信仰："
    get_same_format(document, tittle, faith)

    # 党派
    tittle = "党派："
    get_same_format(document, tittle, partisan)

    # 职业
    tittle = "职业："
    get_same_format(document, tittle, career)

    # 语言
    tittle = "语言："
    get_same_format(document, tittle, language)

    # 证件号码
    tittle = "证件号码："
    judge_text_or_img(document, tittle, id_number)

    # 隶属单位
    tittle = "隶属单位："
    get_same_format(document, tittle, unit)

    # 教育经历
    get_little_tittle(document, "一、教育经历")
    get_big_type(document, ed_experience)

    # 工作经历
    get_little_tittle(document, "二、工作经历")
    get_big_type(document, work_experience)

    # 特点特征
    get_little_tittle(document, "三、特点特征")
    get_big_type(document, characterisics)

    # 经济状况
    get_little_tittle(document, "四、经济状况")
    get_big_type(document, economic_status)

    # 政治主张
    get_little_tittle(document, "五、政治主张")
    get_big_type(document, polotical_views)

    # 犯罪记录
    get_little_tittle(document, "六、犯罪记录")
    get_big_type(document, criminal_record)

    # 争议事件
    get_little_tittle(document, "七、争议事件")
    get_big_type(document, controversial)

    # 重要活动与言论
    get_little_tittle(document, "八、重要活动与言论")
    get_big_type(document, import_event)
    # 主要成就
    get_little_tittle(document, "九、主要成就")
    get_big_type(document, main_achieve)
    # 家庭成员
    get_little_tittle(document, "十、家庭成员")
    get_big_type(document, family_member)
    # 社会关系
    get_little_tittle(document, "十一、社会关系")
    get_big_type(document, social_relationship)
    # 联系方式
    get_little_tittle(document, "十二、联系方式")
    get_big_type(document, contact_information)
    # 其他
    get_little_tittle(document, "十三、其他")
    get_big_type(document, others)

    # document.save('{}.docx'.format(person_name+'调研3'))

    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    filename = quote(person_name + '申请表.doc')
    rv = send_file(f, as_attachment=True, attachment_filename=filename)
    rv.headers['Content-Disposition'] += "; filename*=utf-8''{}".format(filename)
    return rv


# 获取机构的文档
def get_org_docx(id):
    results = get_org_info(id)

    results_dict = results[0]

    # 一、基本情况
    # 组织名称：（中文名称、外文名称）
    association_name = results_dict['association_name']
    # 创建时间：
    create_time = results_dict['create_time']
    # 徽标：
    logo = results_dict['logo']
    # 外观：
    exterior_picture = results_dict['exterior_picture']
    # 宣传口号： 暂无
    #    association_name = results_dict['association_name']
    # 宗旨目的：
    purpose = results_dict['purpose']
    # 负责人：
    leader = results_dict['leader']
    # 上级单位：
    superior_unit = results_dict['superior_unit']
    # 政治主张：
    political_views = results_dict['political_views']

    # 二、组织架构
    organization = results_dict['organization']
    # 三、组织成员
    organization_member = results_dict['organization_member']
    # 四、发展脉络
    development_content = results_dict['development_content']
    # 五、重要活动
    import_event = results_dict['import_event']
    # 六、资金情况
    funding_situation = results_dict['funding_situation']
    # 七、主要成就
    main_achieve = results_dict['main_achieve']
    # 八、负面信息
    negative_information = results_dict['negative_information']
    # 九、关联机构
    associated_institution = results_dict['associated_institution']
    # 十、联系方式
    contact_information = results_dict['contact_information']
    # 十一、其它
    others = results_dict['others']

    # 主文档
    document = Document()

    # 标题
    p = document.add_paragraph()
    run = p.add_run(association_name + "调研")
    run.font.name = u'宋体（正文）'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体（正文）')
    run.font.size = Pt(22)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中
    run.bold = True  # 加粗

    # 简介
    get_little_tittle(document, "一、基本情况")

    # 组织名称：
    tittle = "组织名称："
    get_same_format(document, tittle, association_name)

    # 创建时间：
    tittle = "创建时间："
    get_same_format(document, tittle, create_time)

    # 徽标：
    tittle = "徽标："
    judge_text_or_img(document, tittle, logo)

    # 外观：
    tittle = "外观："
    judge_text_or_img(document, tittle, exterior_picture)
    # 宣传口号： 暂无
    # association_name = results_dict['association_name']

    # 宗旨目的：
    tittle = "宗旨目的："
    judge_text_or_img(document, tittle, purpose)
    # 负责人：
    tittle = "负责人："
    get_same_format(document, tittle, leader)

    # 上级单位：
    tittle = "上级单位："
    get_same_format(document, tittle, superior_unit)

    # 政治主张：
    tittle = "政治主张："
    judge_text_or_img(document, tittle, political_views)

    # 组织架构
    get_little_tittle(document, "二、组织架构")
    get_big_type(document, organization)
    # 三、组织成员
    get_little_tittle(document, "三、组织成员")
    get_big_type(document, organization_member)
    # 四、发展脉络
    get_little_tittle(document, "四、发展脉络")
    get_big_type(document, development_content)
    # 五、重要活动
    get_little_tittle(document, "五、重要活动")
    get_big_type(document, import_event)
    # 六、资金情况
    get_little_tittle(document, "六、资金情况")
    get_big_type(document, funding_situation)
    # 七、主要成就
    get_little_tittle(document, "七、主要成就")
    get_big_type(document, main_achieve)
    # 八、负面信息
    get_little_tittle(document, "八、负面信息")
    get_big_type(document, negative_information)
    # 九、关联机构
    get_little_tittle(document, "九、关联机构")
    get_big_type(document, associated_institution)
    # 十、联系方式
    get_little_tittle(document, "十、联系方式")
    get_big_type(document, contact_information)
    # 十一、其它
    get_little_tittle(document, "十一、其它")
    get_big_type(document, others)


    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    filename = quote(association_name + '分析结果.doc')
    rv = send_file(f, as_attachment=True, attachment_filename=filename)
    rv.headers['Content-Disposition'] += "; filename*=utf-8''{}".format(filename)
    return rv


# 获取事件的文档
def get_event_docx(id):
    results = get_event_info(id)
    results_dict = results[0]
    # 一、事件名字
    event_name = results_dict['event_name']
    # 二、时间
    event_time = results_dict['event_time']
    # 三、地点
    event_address = results_dict['event_address']
    # 四、起因
    event_reason = results_dict['event_reason']
    # 五、经过
    process = results_dict['process']
    # 六、结果
    result = results_dict['result']
    # 七、相关人员
    related_person = results_dict['related_person']
    # 八、关联机构
    associated_institution = results_dict['associated_institution']
    # 九、各方反应
    all_parties = results_dict['all_parties']
    # 十、其它
    others = results_dict['others']


    # 主文档
    document = Document()

    # 标题
    p = document.add_paragraph()
    run = p.add_run(event_name)
    run.font.name = u'宋体（正文）'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体（正文）')
    run.font.size = Pt(22)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中
    run.bold = True  # 加粗


    # 一、事件名字
    get_little_tittle(document, "一、事件名字")
    get_big_type(document, event_name)
    # 二、时间
    get_little_tittle(document, "二、时间")
    get_big_type(document, event_time)
    # 三、地点
    get_little_tittle(document, "三、地点")
    get_big_type(document, event_address)
    # 四、起因
    get_little_tittle(document, "四、起因")
    get_big_type(document, event_reason)
    # 五、经过
    get_little_tittle(document, "五、经过")
    get_big_type(document, process)
    # 六、结果
    get_little_tittle(document, "六、结果")
    get_big_type(document, result)
    # 七、相关人员
    get_little_tittle(document, "七、相关人员")
    get_big_type(document, related_person)
    # 八、关联机构
    get_little_tittle(document, "八、关联机构")
    get_big_type(document, associated_institution)
    # 九、各方反应
    get_little_tittle(document, "九、各方反应")
    get_big_type(document, all_parties)
    # 十、其它
    get_little_tittle(document, "十、其它")
    get_big_type(document, others)

    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    filename = quote(event_name + '.doc')
    rv = send_file(f, as_attachment=True, attachment_filename=filename)
    rv.headers['Content-Disposition'] += "; filename*=utf-8''{}".format(filename)
    return rv


# 获取国别信息的文档
def get_country_docx(id):
    results = get_country_info(id)
    results_dict = results[0]
    # 一、国家名称
    country_name = results_dict['country_name']
    # 二、国旗
    country_img = results_dict['country_img']
    # 三、首都
    capital = results_dict['capital']
    # 四、人口
    population = results_dict['population']
    # 五、货币
    currency = results_dict['currency']
    # 六、官方语言
    official_language = results_dict['official_language']
    # 七、领导人
    leader = results_dict['leader']
    # 八、地图
    map_picture = results_dict['map_picture']
    # 九、地理
    geography = results_dict['geography']
    # 十、历史
    history = results_dict['history']
    # 十一、风俗禁忌
    custom_taboo = results_dict['custom_taboo']
    # 十二、政治
    politics = results_dict['politics']
    # 十三、外交
    diplomatic = results_dict['diplomatic']
    # 十四、军事
    military = results_dict['military']
    # 十五、经济
    economic = results_dict['economic']
    # 十六、国家安全
    national_security = results_dict['national_security']
    # 十七、对华态度
    attitude_china = results_dict['attitude_china']
    # 十八、社会问题
    social_issues = results_dict['social_issues']
    # 十九、建设与环境
    environment = results_dict['environment']
    # 二十、其它
    others = results_dict['others']


    # 主文档
    document = Document()

    # 标题
    p = document.add_paragraph()
    run = p.add_run(country_name)
    run.font.name = u'宋体（正文）'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体（正文）')
    run.font.size = Pt(22)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中
    run.bold = True  # 加粗

    # 一、国家名称
    get_little_tittle(document, "一、国家名称")
    get_big_type(document, country_name)
    # 二、国旗
    get_little_tittle(document, "二、国旗")
    get_big_type(document, country_img)
    # 三、首都
    get_little_tittle(document, "三、首都")
    get_big_type(document, capital)
    # 四、人口
    get_little_tittle(document, "四、人口")
    get_big_type(document, population)
    # 五、货币
    get_little_tittle(document, "五、货币")
    get_big_type(document, currency)
    # 六、官方语言
    get_little_tittle(document, "六、官方语言")
    get_big_type(document, official_language)
    # 七、领导人
    get_little_tittle(document, "七、领导人")
    get_big_type(document, leader)
    # 八、地图
    get_little_tittle(document, "八、地图")
    get_big_type(document, map_picture)
    # 九、地理
    get_little_tittle(document, "九、地理")
    get_big_type(document, geography)
    # 十、历史
    get_little_tittle(document, "十、历史")
    get_big_type(document, history)
    # 十一、风俗禁忌
    get_little_tittle(document, "十一、风俗禁忌")
    get_big_type(document, custom_taboo)
    # 十二、政治
    get_little_tittle(document, "十二、政治")
    get_big_type(document, politics)
    # 十三、外交
    get_little_tittle(document, "十三、外交")
    get_big_type(document, diplomatic)
    # 十四、军事
    get_little_tittle(document, "十四、军事")
    get_big_type(document, military)
    # 十五、经济
    get_little_tittle(document, "十五、经济")
    get_big_type(document, economic)
    # 十六、国家安全
    get_little_tittle(document, "十六、国家安全")
    get_big_type(document, national_security)
    # 十七、对华态度
    get_little_tittle(document, "十七、对华态度")
    get_big_type(document, attitude_china)
    # 十八、社会问题
    get_little_tittle(document, "十八、社会问题")
    get_big_type(document, social_issues)
    # 十九、建设与环境
    get_little_tittle(document, "十九、建设与环境")
    get_big_type(document, environment)
    # 二十、其它
    get_little_tittle(document, "二十、其它")
    get_big_type(document, others)

    # 直接写入内存
    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    filename = quote(country_name + '.doc')
    rv = send_file(f, as_attachment=True, attachment_filename=filename)
    rv.headers['Content-Disposition'] += "; filename*=utf-8''{}".format(filename)
    return rv


# views.py中
@web.route('/download_word', methods=['POST'])
def dowload_word():
    data = request.form.get('data', '')
    if data:
        data = json.loads(data)
        print(data)
        id = data['entity_id']
        if data['entity_type'] == "person":

            result = get_person_docx(id)
        elif data['entity_type'] == "org":
            result = get_org_docx(id)

        elif data['entity_type'] == "event":
            result = get_event_docx(id)

        elif data['entity_type'] == "country":
            result = get_country_docx(id)

        else:
            return jsonify({"message": "no this entity_type"})
        if result:
            return result

        return jsonify({"message": "have error"}), 406
    else:
        return jsonify({"message": "Please pass in the data field ..."}), 400

# end...